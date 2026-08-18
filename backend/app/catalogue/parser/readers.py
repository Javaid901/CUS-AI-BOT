"""
backend/app/catalogue/parser/readers.py

File-format readers for curriculum document uploads.

Each reader turns a file on disk into a uniform structure:

    DocumentText:
      pages: list[{page, text}]        — page-like text sections (for text formats)
      tables: list[list[list[str]]]    — tabular data (spreadsheets / CSV) as
                                         row-major cell grids, header row included
      text: str                        — full combined text (best-effort)
      warnings: list[str]              — non-fatal extraction issues

Formats:
  pdf, docx  -> reuse the existing document pipeline (app.utils.files.extract_text)
  doc        -> lightweight legacy-OLE best-effort reader (spaces/printable)
  csv, xlsx, xls -> tabular readers that also build a readable text form

Future-proofing: add a new format by registering an entry in FORMAT_READERS
(reader signature: `def read(path: str) -> DocumentText`) — no other code
needs to change.
"""

from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass, field
from pathlib import Path

CUR_EXTENSIONS: tuple[str, ...] = ("pdf", "docx", "doc", "csv", "xlsx", "xls")
CUR_EXTENSION_LABEL = ", ".join(CUR_EXTENSIONS)


@dataclass
class DocumentText:
    pages: list[dict] = field(default_factory=list)          # {"page": int, "text": str}
    tables: list[list[list]] = field(default_factory=list)   # row grids (spreadsheets/CSV)
    warnings: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join((p.get("text") or "") for p in self.pages)


class FormatNotSupportedError(Exception):
    pass


# ---------------------------------------------------------------------------
# Text readers
# ---------------------------------------------------------------------------


def _read_pdf(path: str) -> DocumentText:
    from app.utils.files import extract_text

    pages = extract_text(path, "pdf")  # list of {"page", "text"}
    return DocumentText(pages=pages)


def _read_docx(path: str) -> DocumentText:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Prefer table rows too — curriculums hide subject grids in docx tables.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                paras.append("\t".join(cells))
    return DocumentText(pages=[{"page": 1, "text": "\n".join(paras)}])


def _read_doc(path: str) -> DocumentText:
    """Legacy binary .doc — best-effort printable-text extraction.

    Old Word .doc files are OLE compound documents; without a full binary
    parser only the printable/embedded text can be recovered. A warning is
    emitted when little readable text is present so admins know to prefer
    DOCX/PDF uploads.
    """
    raw = Path(path).read_bytes()
    text = re.sub(rb"[^\x20-\x7E\r\n\t]", b" ", raw).decode("ascii", errors="ignore")
    text = re.sub(r"[ \t]{2,}", " ", text)
    lines = [ln.strip() for ln in text.splitlines() if len(ln.strip()) > 1]
    joined = "\n".join(lines)
    warnings = []
    if len(joined) < 200:
        warnings.append(
            "Limited text could be recovered from this legacy DOC file. "
            "Uploading a PDF/DOCX version produces a much better extraction."
        )
    return DocumentText(
        pages=[{"page": 1, "text": joined}],
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# Tabular readers (spreadsheets / CSV)
# ---------------------------------------------------------------------------


def _normalise_cells(row) -> list:
    out = []
    for cell in row:
        if cell is None:
            out.append("")
            continue
        if isinstance(cell, float):
            if cell != cell or str(float(cell)).lower() == "nan":  # NaN guard
                out.append("")
                continue
            if cell.is_integer():
                cell = int(cell)
        out.append(str(cell).strip())
    return out


def _read_xlsx(path: str) -> DocumentText:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    tables: list[list[list]] = []
    rows_txt: list[str] = []
    for ws in wb.worksheets:
        grid: list[list] = []
        for row in ws.iter_rows(values_only=True):
            cells = _normalise_cells(row)
            if cells:
                grid.append(cells)
                rows_txt.append("\t".join(cells))
        if grid:
            tables.append(grid)
    return DocumentText(
        pages=[{"page": 1, "text": "\n".join(rows_txt)}],
        tables=tables,
    )


def _read_xls(path: str) -> DocumentText:
    import xlrd

    book = xlrd.open_workbook(path)
    tables: list[list[list]] = []
    rows_txt: list[str] = []
    for sheet in book.sheets():
        grid: list[list] = []
        for r in range(sheet.nrows):
            cells = _normalise_cells(sheet.row_values(r))
            if cells:
                grid.append(cells)
                rows_txt.append("\t".join(cells))
        if grid:
            tables.append(grid)
    return DocumentText(
        pages=[{"page": 1, "text": "\n".join(rows_txt)}],
        tables=tables,
    )


def _read_csv(path: str) -> DocumentText:
    with open(path, "r", encoding="utf-8-sig", errors="replace", newline="") as fh:
        sample = fh.read(131072)  # sniff dialect from first 128KB
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
    except csv.Error:
        dialect = csv.excel
    rows: list[list] = []
    rows_txt: list[str] = []
    with open(path, "r", encoding="utf-8-sig", errors="replace", newline="") as fh:
        reader = csv.reader(fh, dialect)
        for line in reader:
            cells = [c.strip() for c in line]
            if cells:
                rows.append(cells)
                rows_txt.append("\t".join(cells))
    return DocumentText(
        pages=[{"page": 1, "text": "\n".join(rows_txt)}],
        tables=[rows],
    )


# ---------------------------------------------------------------------------
# Registry (future-proof: add a format by inserting a new reader here)
# ---------------------------------------------------------------------------

FORMAT_READERS: dict[str, callable] = {
    "pdf": _read_pdf,
    "docx": _read_docx,
    "doc": _read_doc,
    "csv": _read_csv,
    "xlsx": _read_xlsx,
    "xls": _read_xls,
}


def supported_extensions() -> list[str]:
    return list(FORMAT_READERS.keys())


def read_curriculum_document(path: str) -> DocumentText:
    """Route a curriculum upload to its format reader (raises FormatNotSupportedError)."""
    ext = Path(path).suffix.lstrip(".").lower()
    reader = FORMAT_READERS.get(ext)
    if reader is None:
        raise FormatNotSupportedError(f"Unsupported curriculum format: {ext}")
    try:
        return reader(path)
    except ImportError as exc:
        raise FormatNotSupportedError(f"Reader library for '.{ext}' is not installed") from exc